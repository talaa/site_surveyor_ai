# tools.py
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageEnhance
import io
import numpy as np
import base64
import re
import os
import io
import tempfile
import shutil
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime
    


TEMP_DIR = "temp_images"  # Temporary directory for images

# Function to extract JSON from markdown
def extract_json(output: str) -> str:
    match = re.search(r'```json\n(.*?)\n```', output, re.DOTALL)
    return match.group(1) if match else output

def file_to_base64(uploaded_file):
    #file_content = uploaded_file.read()  # Read the file’s raw bytes
    base64_string = base64.b64encode(uploaded_file).decode('utf-8')  # Convert to base64
    return base64_string

def enhance_image_quality(image_bytes):
    """
    Enhance the quality of the image for better analysis.
    This is a simplified version that could be expanded with actual image processing.
    
    Args:
        image_bytes: The raw image bytes
        
    Returns:
        bytes: The enhanced image bytes
    """
    try:
        # Open the image using PIL
        img = Image.open(io.BytesIO(image_bytes))
        
        # Example enhancements (adjust based on your needs):
        # 1. Resize if too large (for faster processing)
        max_size = 1500
        if max(img.size) > max_size:
            ratio = max_size / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, Image.LANCZOS)
        
        # 2. Convert to RGB if needed
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # 3. Optional: Adjust contrast, brightness, etc.
        # from PIL import ImageEnhance
        # enhancer = ImageEnhance.Contrast(img)
        # img = enhancer.enhance(1.2)  # Increase contrast by 20%
        
        # Save the enhanced image to bytes
        output = io.BytesIO()
        img.save(output, format=img.format or 'JPEG', quality=95)
        output.seek(0)
        
        return output.getvalue()
        
    except Exception as e:
        print(f"Error enhancing image: {e}")
        # Return original if enhancement fails
        return image_bytes

def generate_report(site_data):
    """
    Generate report files from site data, now handling image bytes directly.
    
    Args:
        site_data: Dictionary containing all the site information and image bytes
        
    Returns:
        dict: Dictionary with binary content of report files
    """
    # Create temp directory for any needed file operations
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Create Word Document
        doc = Document()
        doc.add_heading('Telecom Site Survey Report', 0)
        
        # Create Excel workbook
        excel_data = {}
        
        # Process each section
        for section_name, items in site_data["sections"].items():
            # Word document section
            doc.add_heading(section_name, level=1)
            
            # Excel sheet data
            excel_data[section_name] = []
            
            for item in items:
                # Word document content
                doc.add_heading(f"Image: {item['filename']}", level=2)
                
                # Add image to document
                if 'image_bytes' in item:
                    # Create a temporary file for the image
                    img_path = os.path.join(temp_dir, item['filename'])
                    with open(img_path, 'wb') as img_file:
                        img_file.write(item['image_bytes'])
                    
                    # Add to Word doc
                    doc.add_picture(img_path, width=Inches(6))
                
                # Add analysis text
                doc.add_heading('Analysis', level=3)
                doc.add_paragraph(item['content'])
                doc.add_page_break()
                
                # Collect data for Excel
                excel_data[section_name].append({
                    'Filename': item['filename'],
                    'Analysis': item['content']
                })
        
        # Save Word document to memory
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        """
        # Create Excel file with data from all sections
        xlsx_bytes = io.BytesIO()
        with pd.ExcelWriter(xlsx_bytes, engine='xlsxwriter') as writer:
            # Summary sheet
            summary_data = []
            for section, items in site_data["sections"].items():
                summary_data.append({
                    'Category': section,
                    'Number of Images': len(items)
                })
            
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Section sheets
            for section, items in excel_data.items():
                if items:  # Only create sheets for sections with data
                    df = pd.DataFrame(items)
                    df.to_excel(writer, sheet_name=section[:31], index=False)  # Excel sheet names limited to 31 chars
        
        xlsx_bytes.seek(0)
        """
        return {
            "docx": docx_bytes.getvalue()
            #"xlsx": xlsx_bytes.getvalue()
        }
        
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir)


def create_word_report(data: dict) -> bytes:
    """Generate Word document"""
    doc = Document()
    doc.add_heading('Telecom Site Survey Report', 0)
    
    for section, items in data["sections"].items():
        doc.add_heading(section, level=1)
        for item in items:
            doc.add_paragraph(f"File: {item['filename']}")
            for key, value in item.items():
                if key != 'filename':
                    doc.add_paragraph(f"{key}: {value}")
            doc.add_page_break()
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
"""
def create_excel_report(data: dict) -> bytes:
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Site Analysis"
    
    # Add headers
    headers = ["Section", "Filename", "Details"]
    ws.append(headers)
    
    # Add data
    for section, items in data["sections"].items():
        for item in items:
            ws.append([
                section,
                item['filename'],
                str({k: v for k, v in item.items() if k != 'filename'})
            ])
    
    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()
"""

def clean_temp_folder():
    """Clean any temporary files that might have been created."""
    temp_dir = os.path.join(os.getcwd(), 'tmp')
    if os.path.exists(temp_dir):
        for file in os.listdir(temp_dir):
            file_path = os.path.join(temp_dir, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                print(f"Error deleting {file_path}: {e}")
"""
def save_uploaded_file(uploaded_file):
   
    os.makedirs(TEMP_DIR, exist_ok=True)  # Ensure temp folder exists
    file_path = os.path.join(TEMP_DIR, uploaded_file.name)

    with open(file_path, "wb") as f:
        f.write(uploaded_file.read())  # Save uploaded image

    return file_path
"""